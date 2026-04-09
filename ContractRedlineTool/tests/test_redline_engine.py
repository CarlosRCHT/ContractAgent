import os
import pytest
from docx import Document
from lxml import etree
from app.services.redline_engine import RedlineEngine

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def qn(tag):
    return f"{{{WORD_NS}}}{tag}"


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample Word document for testing."""
    doc = Document()
    doc.add_paragraph("This contract contains unlimited liability provisions.")
    doc.add_paragraph("The auto-renewal period shall be 5 years from the effective date.")
    doc.add_paragraph("Payment terms are Net 120 days from invoice date.")
    doc.add_paragraph("This agreement shall be governed by the laws of the applicable jurisdiction.")
    path = str(tmp_path / "test_contract.docx")
    doc.save(path)
    return path


@pytest.fixture
def multi_paragraph_docx(tmp_path):
    """Create a document with varied content for edge case testing."""
    doc = Document()
    doc.add_heading("VENDOR SERVICES AGREEMENT", level=1)
    doc.add_paragraph("Between Contoso and GlobalServe Consulting Ltd.")
    doc.add_paragraph("")  # Empty paragraph
    doc.add_paragraph("Section 5. LIMITATION OF LIABILITY")
    doc.add_paragraph(
        "The Vendor shall bear unlimited liability for all damages, "
        "losses, and claims arising from or related to the services."
    )
    doc.add_paragraph("Section 6. PAYMENT TERMS")
    doc.add_paragraph("All invoices shall be payable Net 120 days from date of receipt.")
    doc.add_paragraph("Section 7. TERM AND RENEWAL")
    doc.add_paragraph(
        "This Agreement shall automatically renew for successive 5-year periods "
        "unless terminated by either party with 30 days notice."
    )
    path = str(tmp_path / "multi_test.docx")
    doc.save(path)
    return path


class TestRedlineEngineInit:
    def test_init_valid_docx(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        assert engine.doc is not None

    def test_init_invalid_file(self, tmp_path):
        bad_path = str(tmp_path / "nonexistent.docx")
        with pytest.raises(Exception):
            RedlineEngine(bad_path)

    def test_init_creates_comments_part(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        assert engine._comments_element is not None


class TestTextSearch:
    def test_find_existing_text(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        results = engine._find_text_in_paragraphs("unlimited liability")
        assert len(results) >= 1

    def test_find_case_insensitive(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        results = engine._find_text_in_paragraphs("UNLIMITED LIABILITY")
        assert len(results) >= 1

    def test_find_nonexistent_text(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        results = engine._find_text_in_paragraphs("text that does not exist anywhere")
        assert len(results) == 0


class TestApplyTrackedChange:
    def test_apply_single_change(self, sample_docx, tmp_path):
        engine = RedlineEngine(sample_docx)
        result = engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="liability capped at total contract value",
            author="Test Agent",
            rationale="Prohibited per policy §4.2",
            risk_level="major",
        )
        assert result.applied is True
        assert result.comment_added is True
        assert result.error == ""

        # Save and verify the output is a valid docx
        out = str(tmp_path / "output.docx")
        engine.save(out)
        assert os.path.exists(out)

        # Verify the output can be opened
        doc = Document(out)
        assert len(doc.paragraphs) > 0

    def test_apply_change_creates_del_element(self, sample_docx, tmp_path):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="liability capped at total contract value",
            author="Test Agent",
            rationale="Prohibited term",
        )
        out = str(tmp_path / "output.docx")
        engine.save(out)

        # Parse the XML and check for w:del and w:ins
        doc = Document(out)
        body_xml = doc.element.body.xml
        assert "w:del" in body_xml
        assert "w:ins" in body_xml
        assert "w:delText" in body_xml

    def test_apply_change_creates_comment(self, sample_docx, tmp_path):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="capped liability",
            author="Test Agent",
            rationale="Prohibited per policy",
            risk_level="major",
        )
        out = str(tmp_path / "output.docx")
        engine.save(out)

        # Check comment markers in body XML
        doc = Document(out)
        body_xml = doc.element.body.xml
        assert "w:commentRangeStart" in body_xml
        assert "w:commentRangeEnd" in body_xml
        assert "w:commentReference" in body_xml

    def test_apply_change_text_not_found(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        result = engine.apply_tracked_change(
            original_text="nonexistent text in document",
            replacement_text="replacement",
            author="Test Agent",
            rationale="Test",
        )
        assert result.applied is False
        assert "not found" in result.error.lower()

    def test_apply_deletion_only(self, sample_docx, tmp_path):
        """Test deleting text without replacement."""
        engine = RedlineEngine(sample_docx)
        result = engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="",
            author="Test Agent",
            rationale="Remove prohibited term",
            risk_level="major",
        )
        assert result.applied is True
        out = str(tmp_path / "output.docx")
        engine.save(out)

        doc = Document(out)
        body_xml = doc.element.body.xml
        assert "w:del" in body_xml


class TestApplyAllChanges:
    def test_apply_multiple_changes(self, multi_paragraph_docx, tmp_path):
        engine = RedlineEngine(multi_paragraph_docx)
        recommendations = [
            {
                "original_text": "unlimited liability",
                "replacement_text": "liability capped at total contract value",
                "rationale": "Prohibited per policy §4.2",
                "risk_level": "major",
            },
            {
                "original_text": "Net 120 days",
                "replacement_text": "Net 30 days",
                "rationale": "Payment terms exceed policy maximum of Net 60",
                "risk_level": "moderate",
            },
            {
                "original_text": "5-year periods",
                "replacement_text": "1-year periods",
                "rationale": "Auto-renewal exceeds 3-year policy limit",
                "risk_level": "major",
            },
        ]

        results = engine.apply_all_changes(recommendations, author="Test Agent")
        assert len(results) == 3
        applied = sum(1 for r in results if r.applied)
        assert applied == 3

        out = str(tmp_path / "multi_output.docx")
        engine.save(out)
        assert os.path.exists(out)

    def test_apply_mix_of_found_and_not_found(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        recommendations = [
            {
                "original_text": "unlimited liability",
                "replacement_text": "capped liability",
                "rationale": "Policy violation",
                "risk_level": "major",
            },
            {
                "original_text": "nonexistent clause text",
                "replacement_text": "replacement",
                "rationale": "N/A",
                "risk_level": "minor",
            },
        ]

        results = engine.apply_all_changes(recommendations, author="Test Agent")
        assert results[0].applied is True
        assert results[1].applied is False


class TestRiskLevelInComments:
    def test_major_risk_emoji(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="capped liability",
            author="Test Agent",
            rationale="Policy violation",
            risk_level="major",
        )
        comments_xml = etree.tostring(engine._comments_element, encoding="unicode")
        assert "MAJOR" in comments_xml

    def test_moderate_risk_emoji(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="Net 120 days",
            replacement_text="Net 30 days",
            author="Test Agent",
            rationale="Exceeds policy maximum",
            risk_level="moderate",
        )
        comments_xml = etree.tostring(engine._comments_element, encoding="unicode")
        assert "MODERATE" in comments_xml

    def test_minor_risk_emoji(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="the applicable jurisdiction",
            replacement_text="the applicable jurisdiction, as determined by local law",
            author="Test Agent",
            rationale="Minor clarification",
            risk_level="minor",
        )
        comments_xml = etree.tostring(engine._comments_element, encoding="unicode")
        assert "MINOR" in comments_xml


class TestSaveDocument:
    def test_save_default_path(self, sample_docx):
        engine = RedlineEngine(sample_docx)
        engine.apply_tracked_change(
            original_text="unlimited liability",
            replacement_text="capped liability",
            author="Test Agent",
            rationale="Test",
        )
        result_path = engine.save()
        assert os.path.exists(result_path)
        assert "_redlined" in result_path
        os.remove(result_path)  # cleanup

    def test_save_custom_path(self, sample_docx, tmp_path):
        engine = RedlineEngine(sample_docx)
        custom_path = str(tmp_path / "custom_output.docx")
        engine.save(custom_path)
        assert os.path.exists(custom_path)
