import os
import pytest
from unittest.mock import patch, AsyncMock, MagicMock
from fastapi.testclient import TestClient
from docx import Document
from app.main import app


@pytest.fixture
def client():
    return TestClient(app)


@pytest.fixture(autouse=True)
def disable_api_key(monkeypatch):
    """Disable API key authentication for tests."""
    monkeypatch.setattr("app.config.settings.api_key", "")


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample Word document."""
    doc = Document()
    doc.add_paragraph("This contract contains unlimited liability provisions.")
    doc.add_paragraph("Payment terms are Net 120 days.")
    path = str(tmp_path / "test.docx")
    doc.save(path)
    return path


@pytest.fixture
def redline_request():
    return {
        "document_url": "https://contoso.sharepoint.com/sites/contracts/Shared Documents/test.docx",
        "recommendations": [
            {
                "original_text": "unlimited liability",
                "replacement_text": "liability capped at contract value",
                "rationale": "Prohibited per policy §4.2",
                "risk_level": "major",
                "section": "Limitation of Liability",
            }
        ],
        "author": "Test Agent",
    }


@pytest.fixture(autouse=True)
def mock_temp_dir(tmp_path, monkeypatch):
    """Ensure temp_dir exists and is writable on any OS."""
    temp = str(tmp_path / "temp")
    os.makedirs(temp, exist_ok=True)
    monkeypatch.setattr("app.config.settings.temp_dir", temp)


class TestHealthEndpoint:
    def test_health_returns_200(self, client):
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "healthy"
        assert data["version"] == "1.0.0"


class TestRedlineEndpoint:
    @patch("app.routers.redline.graph_client")
    def test_successful_redline(self, mock_graph, client, sample_docx, redline_request):
        mock_graph.download_file = AsyncMock(return_value=sample_docx)
        mock_graph.upload_file = AsyncMock(
            return_value="https://contoso.sharepoint.com/sites/contracts/Shared Documents/test_redlined.docx"
        )
        mock_graph.is_sharing_link = MagicMock(return_value=False)
        mock_graph.parse_sharepoint_url = MagicMock(return_value={
            "hostname": "contoso.sharepoint.com",
            "site_path": "/sites/contracts",
            "file_path": "/Shared Documents/test.docx",
            "filename": "test.docx",
            "library_name": "Shared Documents",
            "item_path": "/test.docx",
        })
        mock_graph.cleanup_temp_file = MagicMock()

        resp = client.post("/api/redline", json=redline_request)
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert data["changesApplied"] == 1
        assert data["commentsAdded"] == 1
        assert "redlined" in data["outputUrl"]

    @patch("app.routers.redline.graph_client")
    def test_text_not_found_returns_partial(self, mock_graph, client, sample_docx):
        mock_graph.download_file = AsyncMock(return_value=sample_docx)
        mock_graph.upload_file = AsyncMock(return_value="https://example.com/out.docx")
        mock_graph.is_sharing_link = MagicMock(return_value=False)
        mock_graph.parse_sharepoint_url = MagicMock(return_value={
            "hostname": "contoso.sharepoint.com",
            "site_path": "/sites/contracts",
            "file_path": "/Shared Documents/test.docx",
            "filename": "test.docx",
            "library_name": "Shared Documents",
            "item_path": "/test.docx",
        })
        mock_graph.cleanup_temp_file = MagicMock()

        request = {
            "document_url": "https://contoso.sharepoint.com/sites/contracts/Shared Documents/test.docx",
            "recommendations": [
                {
                    "original_text": "unlimited liability",
                    "replacement_text": "capped liability",
                    "rationale": "Policy",
                    "risk_level": "major",
                },
                {
                    "original_text": "text that does not exist",
                    "replacement_text": "replacement",
                    "rationale": "N/A",
                    "risk_level": "minor",
                },
            ],
            "author": "Test Agent",
        }

        resp = client.post("/api/redline", json=request)
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "partial"
        assert data["changesApplied"] == 1
        assert data["changesFailed"] == 1

    @patch("app.routers.redline.graph_client")
    def test_sharepoint_download_failure(self, mock_graph, client, redline_request):
        from app.services.graph_client import GraphClientError

        mock_graph.download_file = AsyncMock(side_effect=GraphClientError("Access denied"))

        resp = client.post("/api/redline", json=redline_request)
        assert resp.status_code == 502

    def test_invalid_request_empty_recommendations(self, client):
        resp = client.post("/api/redline", json={
            "document_url": "https://contoso.sharepoint.com/test.docx",
            "recommendations": [],
        })
        assert resp.status_code == 422

    def test_invalid_request_missing_url(self, client):
        resp = client.post("/api/redline", json={
            "recommendations": [
                {
                    "original_text": "test",
                    "replacement_text": "replacement",
                    "rationale": "test reason",
                }
            ],
        })
        assert resp.status_code == 422


class TestExtractTextEndpoint:
    @patch("app.routers.redline.graph_client")
    def test_successful_extract(self, mock_graph, client, sample_docx):
        mock_graph.download_file = AsyncMock(return_value=sample_docx)
        mock_graph.is_sharing_link = MagicMock(return_value=False)
        mock_graph.parse_sharepoint_url = MagicMock(return_value={
            "hostname": "contoso.sharepoint.com",
            "site_path": "/sites/contracts",
            "file_path": "/Shared Documents/test.docx",
            "filename": "test.docx",
            "library_name": "Shared Documents",
            "item_path": "/test.docx",
        })
        mock_graph.cleanup_temp_file = MagicMock()

        resp = client.post("/api/extract-text", json={
            "documentUrl": "https://contoso.sharepoint.com/sites/contracts/Shared Documents/test.docx",
        })
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert data["filename"] == "test.docx"
        assert "unlimited liability" in data["text"]
        assert data["pageCount"] >= 1

    @patch("app.routers.redline.graph_client")
    def test_extract_accepts_snake_case(self, mock_graph, client, sample_docx):
        mock_graph.download_file = AsyncMock(return_value=sample_docx)
        mock_graph.is_sharing_link = MagicMock(return_value=False)
        mock_graph.parse_sharepoint_url = MagicMock(return_value={
            "hostname": "contoso.sharepoint.com",
            "site_path": "/sites/contracts",
            "file_path": "/Shared Documents/test.docx",
            "filename": "test.docx",
            "library_name": "Shared Documents",
            "item_path": "/test.docx",
        })
        mock_graph.cleanup_temp_file = MagicMock()

        resp = client.post("/api/extract-text", json={
            "document_url": "https://contoso.sharepoint.com/sites/contracts/Shared Documents/test.docx",
        })
        assert resp.status_code == 200
        assert resp.json()["status"] == "success"

    @patch("app.routers.redline.graph_client")
    def test_extract_sharepoint_failure(self, mock_graph, client):
        from app.services.graph_client import GraphClientError
        mock_graph.download_file = AsyncMock(side_effect=GraphClientError("Access denied"))

        resp = client.post("/api/extract-text", json={
            "documentUrl": "https://contoso.sharepoint.com/sites/contracts/Shared Documents/test.docx",
        })
        assert resp.status_code == 502

    def test_extract_missing_url(self, client):
        resp = client.post("/api/extract-text", json={})
        assert resp.status_code == 422
