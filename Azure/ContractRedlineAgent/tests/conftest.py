"""Shared pytest fixtures.

Adds the ``agent/code_files`` directory to ``sys.path`` so the modules can
be imported by their flat names (``redline_core``, ``docx_text``,
``schemas``) — exactly the same way they're imported inside the Foundry
Code Interpreter sandbox.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import pytest
from docx import Document

CODE_FILES_DIR = (
    Path(__file__).resolve().parent.parent / "agent" / "code_files"
)
sys.path.insert(0, str(CODE_FILES_DIR))


@pytest.fixture
def sample_doc_bytes() -> bytes:
    """Build a small in-memory contract-style .docx and return its bytes."""
    from io import BytesIO

    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph(
        "This Agreement is entered into between Contoso Inc. and the Vendor. "
        "The Vendor shall provide services in accordance with the attached "
        "Statement of Work. Payment shall be made within thirty (30) days of "
        "invoice receipt."
    )
    doc.add_paragraph(
        "Either party may terminate this Agreement for convenience upon "
        "ninety (90) days written notice. Termination for cause requires a "
        "thirty (30) day cure period."
    )
    doc.add_paragraph(
        "The Vendor's liability under this Agreement shall not exceed the "
        "fees paid in the twelve (12) months preceding the claim."
    )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_doc_path(tmp_path: Path, sample_doc_bytes: bytes) -> Path:
    p = tmp_path / "sample.docx"
    p.write_bytes(sample_doc_bytes)
    return p
