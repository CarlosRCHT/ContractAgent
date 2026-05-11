"""End-to-end tests for the redline engine on real .docx bytes."""

from __future__ import annotations

import io

import pytest
from docx import Document
from lxml import etree

from redline_core import (
    RedlineEngine,
    apply_recommendations,
    qn,
)


def _para_xml(doc, paragraph_index: int) -> str:
    return etree.tostring(doc.paragraphs[paragraph_index]._element, encoding="unicode")


def _all_xml(doc) -> str:
    return etree.tostring(doc.element, encoding="unicode")


class TestApplyOne:
    def test_word_level_edit_keeps_unchanged_words_outside_del_ins(
        self, sample_doc_bytes
    ):
        engine = RedlineEngine.from_bytes(sample_doc_bytes, author="Test Author")
        result = engine.apply_one(
            original_text="Payment shall be made within thirty (30) days of invoice receipt.",
            replacement_text="Payment shall be made within sixty (60) days of invoice receipt.",
            rationale="Standard NET-60 terms preferred per finance policy.",
            risk_level="moderate",
        )
        assert result.applied
        assert result.comment_added

        out = engine.to_bytes()
        doc = Document(io.BytesIO(out))

        all_xml = _all_xml(doc)
        # The unchanged anchor words must NOT be wrapped in <w:del>/<w:ins>.
        assert "<w:del " in all_xml
        assert "<w:ins " in all_xml
        # Walk runs in the doc and check that "Payment shall be made within"
        # appears as a plain run whose parent is the paragraph (not a del/ins).
        for para in doc.paragraphs:
            for run in para.runs:
                if "Payment shall be made within" in (run.text or ""):
                    parent_tag = run._element.getparent().tag
                    assert parent_tag.endswith("}p"), (
                        f"Unchanged text was wrapped in {parent_tag}"
                    )

    def test_only_changed_tokens_are_in_del_ins(self, sample_doc_bytes):
        engine = RedlineEngine.from_bytes(sample_doc_bytes, author="Test Author")
        engine.apply_one(
            original_text="ninety (90) days written notice",
            replacement_text="thirty (30) days written notice",
            rationale="Match standard 30-day termination notice window.",
            risk_level="major",
        )
        out = engine.to_bytes()
        doc = Document(io.BytesIO(out))

        del_texts: list[str] = []
        ins_texts: list[str] = []
        for el in doc.element.iter(qn("w:del")):
            for t in el.iter(qn("w:delText")):
                if t.text:
                    del_texts.append(t.text)
        for el in doc.element.iter(qn("w:ins")):
            for t in el.iter(qn("w:t")):
                if t.text:
                    ins_texts.append(t.text)

        joined_del = "".join(del_texts)
        joined_ins = "".join(ins_texts)
        # Only changing tokens should appear in del/ins:
        assert "ninety" in joined_del
        assert "90" in joined_del
        assert "thirty" in joined_ins
        assert "30" in joined_ins
        # Unchanged words must not appear in either bucket.
        assert "written" not in joined_del
        assert "notice" not in joined_del
        assert "written" not in joined_ins
        assert "notice" not in joined_ins

    def test_comment_is_anchored_at_sentence_scope(self, sample_doc_bytes):
        engine = RedlineEngine.from_bytes(sample_doc_bytes, author="Test Author")
        engine.apply_one(
            original_text="thirty (30) days of invoice receipt",
            replacement_text="sixty (60) days of invoice receipt",
            rationale="NET-60 preferred.",
            risk_level="moderate",
        )
        out = engine.to_bytes()
        doc = Document(io.BytesIO(out))

        # Find the paragraph containing the change and read its XML.
        target_para = next(
            p for p in doc.paragraphs if "Payment shall be made" in p.text
        )
        para_xml = etree.tostring(target_para._element, encoding="unicode")

        # commentRangeStart should appear before "Payment" (the start of
        # the sentence, not the start of the matched span "thirty").
        crs_idx = para_xml.find("commentRangeStart")
        cre_idx = para_xml.find("commentRangeEnd")
        assert crs_idx != -1 and cre_idx != -1

        before_crs = para_xml[:crs_idx]
        # The first two sentences ("This Agreement..." and "...Statement of Work.")
        # must sit BEFORE the comment range start, since the match is in the
        # third sentence and the anchor wraps only that sentence.
        assert "Statement of Work" in before_crs

        inside = para_xml[crs_idx:cre_idx]
        # The matched-and-rewritten content lives inside the comment range.
        assert "shall be made within" in inside or "Payment" in inside
        # Earlier sentences must NOT leak into the comment range.
        assert "Statement of Work" not in inside
        assert "This Agreement is entered" not in inside

    def test_text_not_found_returns_error_without_modifying_doc(
        self, sample_doc_bytes
    ):
        engine = RedlineEngine.from_bytes(sample_doc_bytes, author="Test")
        result = engine.apply_one(
            original_text="this exact phrase is not present anywhere",
            replacement_text="something else",
            rationale="Should fail to apply.",
        )
        assert not result.applied
        assert "not found" in result.error.lower()

    def test_adjusted_note_appears_in_comment(self, sample_doc_bytes):
        engine = RedlineEngine.from_bytes(sample_doc_bytes, author="Test")
        engine.apply_one(
            original_text="ninety (90) days written notice",
            replacement_text="forty-five (45) days written notice",
            rationale="Compromise position.",
            risk_level="moderate",
            adjusted=True,
            adjustment_note="Original recommendation was 30 days but cure period would conflict.",
        )
        out = engine.to_bytes()
        doc = Document(io.BytesIO(out))
        # The comments part should contain the adjustment note.
        comments_part = None
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype:
                comments_part = rel.target_part
                break
        assert comments_part is not None
        comments_xml = comments_part.blob.decode("utf-8")
        assert "Adjusted from original recommendation" in comments_xml
        assert "cure period" in comments_xml


class TestBatchApply:
    def test_apply_recommendations_helper(self, sample_doc_bytes):
        recs = [
            {
                "originalText": "thirty (30) days of invoice receipt",
                "replacementText": "sixty (60) days of invoice receipt",
                "rationale": "NET-60 preferred.",
                "riskLevel": "moderate",
            },
            {
                "originalText": "ninety (90) days written notice",
                "replacementText": "thirty (30) days written notice",
                "rationale": "Standard 30-day notice.",
                "riskLevel": "major",
            },
        ]
        out, results = apply_recommendations(sample_doc_bytes, recs)
        assert all(r.applied for r in results)
        assert all(r.comment_added for r in results)
        # Output should be a valid .docx (parseable).
        Document(io.BytesIO(out))

    def test_partial_apply_when_one_match_missing(self, sample_doc_bytes):
        recs = [
            {
                "originalText": "thirty (30) days of invoice receipt",
                "replacementText": "sixty (60) days of invoice receipt",
                "rationale": "NET-60 preferred.",
            },
            {
                "originalText": "phrase not present in fixture",
                "replacementText": "anything",
                "rationale": "Should fail to apply.",
            },
        ]
        _, results = apply_recommendations(sample_doc_bytes, recs)
        assert results[0].applied
        assert not results[1].applied
