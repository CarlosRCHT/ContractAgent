"""Word-level tracked-change engine for the Foundry redline agent.

Differs from ``Azure/ContractRedlineTool`` in two important ways:

* **Word-level edits.** Only the tokens that actually change are wrapped in
  ``<w:del>`` / ``<w:ins>``. Unchanged words remain as plain runs. The diff
  is computed with :class:`difflib.SequenceMatcher` over a lossless
  whitespace-preserving tokenizer.

* **Sentence-scoped comments.** The ``<w:commentRangeStart>`` /
  ``<w:commentRangeEnd>`` anchors wrap the entire affected sentence — or
  the whole paragraph when the match crosses sentence boundaries — even
  though the tracked edits themselves cover only the changing words.

The module is intentionally self-contained (only ``python-docx`` + ``lxml``)
so it can be uploaded as a Foundry agent file and imported from inside the
Code Interpreter sandbox.
"""

from __future__ import annotations

import copy
import difflib
import io
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Iterable, Optional

from docx import Document
from lxml import etree

# Local imports — these are sibling modules also uploaded as agent files.
from docx_text import Span, anchor_span  # type: ignore  # noqa: E402

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# OpenXML namespace plumbing
# ---------------------------------------------------------------------------

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NSMAP = {"w": WORD_NS}


def qn(tag: str) -> str:
    """Return a Clark-form qualified name from a ``w:`` prefixed tag."""
    prefix, local = tag.split(":", 1)
    return f"{{{NSMAP[prefix]}}}{local}"


def _xml_space_preserve(element) -> None:
    element.set(f"{{{XML_NS}}}space", "preserve")


# ---------------------------------------------------------------------------
# Public dataclasses (kept independent from pydantic schemas.py so the engine
# can run in environments where pydantic isn't available)
# ---------------------------------------------------------------------------


@dataclass
class EngineChangeResult:
    original_text: str
    replacement_text: str
    final_replacement_text: str = ""
    applied: bool = False
    adjusted: bool = False
    comment_added: bool = False
    risk_level: str = "moderate"
    error: str = ""


@dataclass
class _RunMapEntry:
    """Maps a run inside a paragraph to its character span."""

    index: int
    start: int
    end: int


@dataclass
class _Match:
    paragraph_index: int
    paragraph: object  # python-docx Paragraph
    match: Span
    anchor: Span


# ---------------------------------------------------------------------------
# Tokenizer + diff
# ---------------------------------------------------------------------------

_TOKEN_RE = re.compile(r"\s+|\w+|[^\w\s]", re.UNICODE)


def tokenize(text: str) -> list[str]:
    """Split *text* into whitespace, word, and punctuation tokens losslessly.

    ``"".join(tokenize(text)) == text`` for any input.
    """
    return _TOKEN_RE.findall(text)


@dataclass
class DiffSegment:
    """One contiguous slice of the word-level diff.

    ``kind`` is one of ``"equal"``, ``"delete"``, ``"insert"``, ``"replace"``.
    For ``"replace"`` both ``del_text`` and ``ins_text`` are populated.
    """

    kind: str
    del_text: str = ""
    ins_text: str = ""


def word_diff(original: str, replacement: str) -> list[DiffSegment]:
    """Return diff segments mapping *original* to *replacement* token-by-token.

    Whitespace-only opcodes are collapsed into the surrounding equal/replace
    segment to avoid emitting noisy single-space deletions.
    """
    a = tokenize(original)
    b = tokenize(replacement)

    matcher = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    segments: list[DiffSegment] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        del_text = "".join(a[i1:i2])
        ins_text = "".join(b[j1:j2])
        if tag == "equal":
            if del_text:
                segments.append(DiffSegment("equal", del_text=del_text, ins_text=del_text))
        elif tag == "delete":
            segments.append(DiffSegment("delete", del_text=del_text))
        elif tag == "insert":
            segments.append(DiffSegment("insert", ins_text=ins_text))
        elif tag == "replace":
            segments.append(DiffSegment("replace", del_text=del_text, ins_text=ins_text))

    # Merge whitespace-only deletes/inserts into adjacent equal segments to
    # avoid Word showing a deleted-then-inserted single space when a phrase
    # is reworded inside an otherwise unchanged sentence.
    return _coalesce_whitespace(segments)


def _coalesce_whitespace(segments: list[DiffSegment]) -> list[DiffSegment]:
    """Fold whitespace-only ``delete``/``insert`` segments into adjacent ``equal`` segments.

    Without this, a reword in the middle of a sentence frequently produces a
    deleted-then-inserted single space which Word renders as a noisy edit
    even though the spacing didn't actually change.
    """
    out: list[DiffSegment] = []
    for seg in segments:
        whitespace_only = (
            (seg.kind == "delete" and seg.del_text and seg.del_text.isspace())
            or (seg.kind == "insert" and seg.ins_text and seg.ins_text.isspace())
        )
        if whitespace_only and out and out[-1].kind == "equal":
            text = seg.del_text if seg.kind == "delete" else seg.ins_text
            out[-1].del_text += text
            out[-1].ins_text += text
            continue
        out.append(seg)
    return out


# ---------------------------------------------------------------------------
# Engine
# ---------------------------------------------------------------------------


class RedlineEngine:
    """Apply tracked changes + comments to a Word document.

    Typical usage from inside Code Interpreter::

        engine = RedlineEngine.from_bytes(doc_bytes, author="Contract Review Agent")
        for rec in recommendations:
            engine.apply_one(
                original_text=rec["originalText"],
                replacement_text=rec["finalReplacementText"],
                rationale=rec["rationale"],
                risk_level=rec.get("riskLevel", "moderate"),
                adjusted=rec.get("adjusted", False),
            )
        out_bytes = engine.to_bytes()
    """

    RISK_EMOJI = {"major": "🔴", "moderate": "🟡", "minor": "🟢"}

    def __init__(self, doc, *, author: str = "Contract Review Agent", date_str: str = "") -> None:
        self.doc = doc
        self.author = author
        self.date_str = date_str or datetime.now(timezone.utc).strftime(
            "%Y-%m-%dT%H:%M:%SZ"
        )
        self._comment_id_counter = 0
        self._revision_id_counter = 100
        self._comments_part = None
        self._comments_element = None
        self._ensure_comments_part()

    # ---- Constructors -----------------------------------------------------

    @classmethod
    def from_path(cls, path: str, **kwargs) -> "RedlineEngine":
        return cls(Document(path), **kwargs)

    @classmethod
    def from_bytes(cls, data: bytes, **kwargs) -> "RedlineEngine":
        return cls(Document(io.BytesIO(data)), **kwargs)

    # ---- Saving -----------------------------------------------------------

    def to_bytes(self) -> bytes:
        self._flush_comments()
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

    def save(self, path: str) -> str:
        self._flush_comments()
        self.doc.save(path)
        return path

    # ---- Public apply API -------------------------------------------------

    def apply_one(
        self,
        *,
        original_text: str,
        replacement_text: str,
        rationale: str,
        risk_level: str = "moderate",
        adjusted: bool = False,
        adjustment_note: str = "",
    ) -> EngineChangeResult:
        """Apply a single recommendation; only tokens that change become tracked edits."""
        result = EngineChangeResult(
            original_text=original_text,
            replacement_text=replacement_text,
            final_replacement_text=replacement_text,
            risk_level=risk_level,
            adjusted=adjusted,
        )

        match = self._locate(original_text)
        if match is None:
            result.error = f"Text not found: {original_text[:60]!r}"
            logger.warning(result.error)
            return result

        comment_id = self._next_comment_id()
        comment_text = self._format_comment(rationale, risk_level, adjusted, adjustment_note)
        self._add_comment(comment_id, comment_text)

        try:
            self._rewrite_paragraph(
                match=match,
                replacement_text=replacement_text,
                comment_id=comment_id,
            )
        except Exception as exc:  # pragma: no cover - defensive
            result.error = f"Failed to apply tracked change: {exc}"
            logger.exception(result.error)
            return result

        result.applied = True
        result.comment_added = True
        return result

    # ---- Match locator ----------------------------------------------------

    def _locate(self, search_text: str) -> Optional[_Match]:
        """Locate *search_text* (case-insensitive) within a paragraph."""
        needle = search_text.lower()
        if not needle.strip():
            return None

        for p_idx, paragraph in enumerate(self.doc.paragraphs):
            full_text = paragraph.text
            if not full_text:
                continue
            idx = full_text.lower().find(needle)
            if idx == -1:
                continue
            span = Span(idx, idx + len(search_text))
            anchor = anchor_span(full_text, span.start, span.end)
            return _Match(p_idx, paragraph, span, anchor)

        return None

    # ---- Paragraph rewrite ------------------------------------------------

    def _rewrite_paragraph(
        self,
        *,
        match: _Match,
        replacement_text: str,
        comment_id: int,
    ) -> None:
        """Rebuild a paragraph's run sequence with diff-aware tracked edits."""
        paragraph = match.paragraph
        p_element = paragraph._element
        full_text = paragraph.text

        # 1. Make sure there are clean run boundaries at the four key offsets.
        for offset in sorted({match.anchor.start, match.match.start, match.match.end, match.anchor.end}):
            if 0 < offset < len(full_text):
                self._ensure_run_boundary(paragraph, offset)

        # 2. Re-walk runs and bucket them by region.
        run_map = self._build_run_map(paragraph)
        anchor_runs: list = []  # runs overlapping [anchor.start, anchor.end)
        match_runs: list = []   # runs strictly inside [match.start, match.end)
        match_run_ids: set = set()
        for entry, run in zip(run_map, paragraph.runs):
            if entry.end <= match.anchor.start or entry.start >= match.anchor.end:
                continue
            anchor_runs.append(run)
            if entry.start >= match.match.start and entry.end <= match.match.end:
                match_runs.append(run)
                match_run_ids.add(id(run._element))

        if not match_runs:
            raise RuntimeError("Match span did not resolve to any runs after splitting.")

        # 3. Capture rPr template from the first matched run.
        rpr_template = self._copy_rpr(match_runs[0]._element)

        # 4. Compute word-level diff and build replacement XML elements.
        original_match_text = full_text[match.match.start:match.match.end]
        diff = word_diff(original_match_text, replacement_text)
        new_elements = self._build_diff_elements(diff, rpr_template)

        # 5. Insert new elements at the position of the first matched run.
        first_match_el = match_runs[0]._element
        anchor_position = first_match_el  # will keep moving as we addprevious

        for el in new_elements:
            first_match_el.addprevious(el)

        # 6. Remove the original matched runs.
        for run in match_runs:
            p_element.remove(run._element)

        # 7. Insert comment range start/end + reference.
        #    The anchor should bracket the entire sentence(s).
        comment_range_start = etree.Element(qn("w:commentRangeStart"))
        comment_range_start.set(qn("w:id"), str(comment_id))
        comment_range_end = etree.Element(qn("w:commentRangeEnd"))
        comment_range_end.set(qn("w:id"), str(comment_id))

        # Re-resolve anchor runs after our edits — anchor_runs that weren't
        # part of match_runs are still valid; the new diff elements sit in
        # place of the removed match runs and should also be inside the
        # comment range. Determine the first/last children of p_element that
        # belong to the anchor region.
        first_anchor_el, last_anchor_el = self._resolve_anchor_bounds(
            paragraph=paragraph,
            anchor=match.anchor,
            new_diff_elements=new_elements,
            kept_anchor_runs=[r for r in anchor_runs if id(r._element) not in match_run_ids],
        )

        first_anchor_el.addprevious(comment_range_start)

        comment_ref_run = etree.Element(qn("w:r"))
        rpr = etree.SubElement(comment_ref_run, qn("w:rPr"))
        style = etree.SubElement(rpr, qn("w:rStyle"))
        style.set(qn("w:val"), "CommentReference")
        ref = etree.SubElement(comment_ref_run, qn("w:commentReference"))
        ref.set(qn("w:id"), str(comment_id))

        last_anchor_el.addnext(comment_range_end)
        comment_range_end.addnext(comment_ref_run)

    # ---- Run boundary management -----------------------------------------

    def _build_run_map(self, paragraph) -> list[_RunMapEntry]:
        entries: list[_RunMapEntry] = []
        cursor = 0
        for i, run in enumerate(paragraph.runs):
            text = run.text or ""
            entries.append(_RunMapEntry(i, cursor, cursor + len(text)))
            cursor += len(text)
        return entries

    def _ensure_run_boundary(self, paragraph, char_offset: int) -> None:
        """Split a run if *char_offset* falls in its interior."""
        cursor = 0
        for run in list(paragraph.runs):
            text = run.text or ""
            run_end = cursor + len(text)
            if char_offset == cursor or char_offset == run_end:
                return
            if cursor < char_offset < run_end:
                inner = char_offset - cursor
                before_text = text[:inner]
                after_text = text[inner:]

                # Modify <w:t> in current run to before_text.
                t_elements = run._element.findall(qn("w:t"))
                if not t_elements:
                    return  # Run has no plain text element — leave as-is.
                t_elements[0].text = before_text
                _xml_space_preserve(t_elements[0])
                # If there were multiple <w:t> children, drop the rest;
                # they'll be reconstructed in the after-run.
                for extra in t_elements[1:]:
                    run._element.remove(extra)

                after_run_el = copy.deepcopy(run._element)
                for t in after_run_el.findall(qn("w:t")):
                    after_run_el.remove(t)
                t_after = etree.SubElement(after_run_el, qn("w:t"))
                t_after.text = after_text
                _xml_space_preserve(t_after)
                run._element.addnext(after_run_el)
                return
            cursor = run_end

    # ---- Diff element construction ---------------------------------------

    def _build_diff_elements(
        self, diff: list[DiffSegment], rpr_template
    ) -> list:
        """Translate diff segments into a list of XML elements ready to splice in."""
        elements: list = []
        for seg in diff:
            if seg.kind == "equal":
                elements.append(self._make_run(seg.ins_text, rpr_template))
            elif seg.kind == "delete":
                elements.append(self._make_del(seg.del_text, rpr_template))
            elif seg.kind == "insert":
                elements.append(self._make_ins(seg.ins_text, rpr_template))
            elif seg.kind == "replace":
                elements.append(self._make_del(seg.del_text, rpr_template))
                elements.append(self._make_ins(seg.ins_text, rpr_template))
        return elements

    def _make_run(self, text: str, rpr_template) -> etree._Element:
        run = etree.Element(qn("w:r"))
        if rpr_template is not None:
            run.append(copy.deepcopy(rpr_template))
        t = etree.SubElement(run, qn("w:t"))
        t.text = text
        _xml_space_preserve(t)
        return run

    def _make_del(self, text: str, rpr_template) -> etree._Element:
        rev_id = self._next_revision_id()
        del_el = etree.Element(qn("w:del"))
        del_el.set(qn("w:id"), str(rev_id))
        del_el.set(qn("w:author"), self.author)
        del_el.set(qn("w:date"), self.date_str)

        run = etree.SubElement(del_el, qn("w:r"))
        if rpr_template is not None:
            run.append(copy.deepcopy(rpr_template))
        del_text = etree.SubElement(run, qn("w:delText"))
        del_text.text = text
        _xml_space_preserve(del_text)
        return del_el

    def _make_ins(self, text: str, rpr_template) -> etree._Element:
        rev_id = self._next_revision_id()
        ins_el = etree.Element(qn("w:ins"))
        ins_el.set(qn("w:id"), str(rev_id))
        ins_el.set(qn("w:author"), self.author)
        ins_el.set(qn("w:date"), self.date_str)

        run = etree.SubElement(ins_el, qn("w:r"))
        if rpr_template is not None:
            run.append(copy.deepcopy(rpr_template))
        t = etree.SubElement(run, qn("w:t"))
        t.text = text
        _xml_space_preserve(t)
        return ins_el

    def _copy_rpr(self, run_element) -> Optional[etree._Element]:
        rpr = run_element.find(qn("w:rPr"))
        return copy.deepcopy(rpr) if rpr is not None else None

    # ---- Comment anchor resolution ---------------------------------------

    def _resolve_anchor_bounds(
        self,
        *,
        paragraph,
        anchor: Span,
        new_diff_elements: list,
        kept_anchor_runs: list,
    ):
        """Return the first/last paragraph children that should sit inside the comment range.

        We always include the new diff elements; we extend leftward to any
        kept anchor run that precedes the diff (start of sentence before
        the change) and rightward to any kept anchor run that follows.
        """
        p_element = paragraph._element
        diff_set = {id(e) for e in new_diff_elements}

        # Walk paragraph children in document order.
        children = list(p_element)
        children_with_index = [(i, c) for i, c in enumerate(children)]

        # Find the index range covering: any kept anchor run + any diff element.
        target_ids = {id(r._element) for r in kept_anchor_runs} | diff_set
        idx_in = [i for i, c in children_with_index if id(c) in target_ids]
        if not idx_in:
            # Fallback — wrap just the diff elements.
            return new_diff_elements[0], new_diff_elements[-1]

        first = children[min(idx_in)]
        last = children[max(idx_in)]
        return first, last

    # ---- Comments part ---------------------------------------------------

    def _ensure_comments_part(self) -> None:
        doc_part = self.doc.part
        for rel in doc_part.rels.values():
            if "comments" in rel.reltype:
                self._comments_part = rel.target_part
                self._comments_element = etree.fromstring(self._comments_part.blob)
                for comment in self._comments_element.findall(qn("w:comment")):
                    cid = int(comment.get(qn("w:id"), "0"))
                    self._comment_id_counter = max(self._comment_id_counter, cid)
                return

        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            "/>"
        )
        self._comments_element = etree.fromstring(comments_xml.encode("utf-8"))

        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        comments_part = Part(
            PackURI("/word/comments.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            etree.tostring(
                self._comments_element,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            ),
            doc_part.package,
        )
        doc_part.relate_to(
            comments_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        )
        self._comments_part = comments_part

    def _add_comment(self, comment_id: int, body_text: str) -> None:
        comment_el = etree.SubElement(self._comments_element, qn("w:comment"))
        comment_el.set(qn("w:id"), str(comment_id))
        comment_el.set(qn("w:author"), self.author)
        comment_el.set(qn("w:date"), self.date_str)
        comment_el.set(qn("w:initials"), self.author[:2].upper())

        p = etree.SubElement(comment_el, qn("w:p"))
        pPr = etree.SubElement(p, qn("w:pPr"))
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), "CommentText")

        ref_run = etree.SubElement(p, qn("w:r"))
        ref_rpr = etree.SubElement(ref_run, qn("w:rPr"))
        ref_style = etree.SubElement(ref_rpr, qn("w:rStyle"))
        ref_style.set(qn("w:val"), "CommentReference")
        etree.SubElement(ref_run, qn("w:annotationRef"))

        body_run = etree.SubElement(p, qn("w:r"))
        body_t = etree.SubElement(body_run, qn("w:t"))
        body_t.text = body_text
        _xml_space_preserve(body_t)

    def _flush_comments(self) -> None:
        if self._comments_part is None or self._comments_element is None:
            return
        self._comments_part._blob = etree.tostring(
            self._comments_element,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    # ---- ID generators / formatting --------------------------------------

    def _next_comment_id(self) -> int:
        self._comment_id_counter += 1
        return self._comment_id_counter

    def _next_revision_id(self) -> int:
        self._revision_id_counter += 1
        return self._revision_id_counter

    def _format_comment(
        self, rationale: str, risk_level: str, adjusted: bool, adjustment_note: str
    ) -> str:
        emoji = self.RISK_EMOJI.get(risk_level, "🟡")
        head = f"{emoji} [{risk_level.upper()}] {rationale}"
        if adjusted and adjustment_note:
            head += f"\n\nAdjusted from original recommendation: {adjustment_note}"
        elif adjusted:
            head += "\n\n(Wording adjusted to fit surrounding context.)"
        return head


# ---------------------------------------------------------------------------
# Convenience batch API
# ---------------------------------------------------------------------------


def apply_recommendations(
    doc_bytes: bytes,
    recommendations: Iterable[dict],
    *,
    author: str = "Contract Review Agent",
) -> tuple[bytes, list[EngineChangeResult]]:
    """High-level helper used by the agent's Code Interpreter step.

    ``recommendations`` is an iterable of dicts with keys::

        originalText, replacementText, rationale, riskLevel,
        adjusted (optional bool), adjustmentNote (optional str)
    """
    engine = RedlineEngine.from_bytes(doc_bytes, author=author)
    results: list[EngineChangeResult] = []
    for rec in recommendations:
        results.append(
            engine.apply_one(
                original_text=rec["originalText"],
                replacement_text=rec.get("replacementText", ""),
                rationale=rec["rationale"],
                risk_level=rec.get("riskLevel", "moderate"),
                adjusted=bool(rec.get("adjusted", False)),
                adjustment_note=rec.get("adjustmentNote", ""),
            )
        )
    return engine.to_bytes(), results
