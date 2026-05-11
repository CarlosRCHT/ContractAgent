"""Local end-to-end smoke test — no Foundry, no Graph required.

Builds a sample contract, runs the redline pipeline directly against the
engine (mocking the Logic App boundaries), and writes the redlined output
next to the source so you can open it in Word and inspect the result.

Usage::

    python scripts/smoke_test.py
"""

from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "agent" / "code_files"))

from redline_core import apply_recommendations  # noqa: E402


def build_sample_contract() -> bytes:
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph(
        "This Agreement is entered into between Contoso Inc. and the Vendor. "
        "The Vendor shall provide services in accordance with the attached "
        "Statement of Work. Payment shall be made within thirty (30) days of "
        "invoice receipt."
    )
    doc.add_paragraph(
        "Either party may terminate this Agreement for convenience upon "
        "ninety (90) days written notice. Termination for cause requires a "
        "thirty (30) day cure period."
    )
    doc.add_paragraph(
        "The Vendor's liability under this Agreement shall not exceed the "
        "fees paid in the twelve (12) months preceding the claim."
    )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    out_dir = ROOT / "tests" / "fixtures"
    out_dir.mkdir(parents=True, exist_ok=True)

    src_path = out_dir / "smoke_source.docx"
    redlined_path = out_dir / "smoke_source_redlined.docx"

    src_bytes = build_sample_contract()
    src_path.write_bytes(src_bytes)
    print(f"Wrote source: {src_path}")

    recommendations = [
        {
            "originalText": "thirty (30) days of invoice receipt",
            "replacementText": "sixty (60) days of invoice receipt",
            "rationale": "NET-60 preferred per finance policy.",
            "riskLevel": "moderate",
        },
        {
            "originalText": "ninety (90) days written notice",
            "replacementText": "thirty (30) days written notice",
            "rationale": "Standard 30-day termination notice window.",
            "riskLevel": "major",
            "adjusted": True,
            "adjustmentNote": "Original recommended 14-day notice; extended to 30 to match cure-period clause.",
        },
        {
            "originalText": "fees paid in the twelve (12) months preceding the claim",
            "replacementText": "fees paid in the twenty-four (24) months preceding the claim",
            "rationale": "Cap aligned with broader liability framework.",
            "riskLevel": "minor",
        },
    ]

    out_bytes, results = apply_recommendations(src_bytes, recommendations)
    redlined_path.write_bytes(out_bytes)

    print(f"\nWrote redlined: {redlined_path}\n")
    print(f"{'#':<3} {'applied':<8} {'adjusted':<9} {'risk':<9} text")
    print("-" * 80)
    for i, r in enumerate(results, 1):
        snippet = (r.original_text[:48] + "...") if len(r.original_text) > 50 else r.original_text
        print(
            f"{i:<3} {str(r.applied):<8} {str(r.adjusted):<9} "
            f"{r.risk_level:<9} {snippet}"
        )
        if r.error:
            print(f"     ERROR: {r.error}")


if __name__ == "__main__":
    main()
